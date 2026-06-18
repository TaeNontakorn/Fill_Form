from back_end import wrap_values, _MARKER, process_runs_to_underline
from docx import Document

data = {
    'a': 'Hello',
    'b': 123,
    'c': 45.6,
    'd': False,
    'e': None,
    'f': ['x', 7]
}
wrapped = wrap_values(data)
print('wrapped=', wrapped)

doc = Document()
p = doc.add_paragraph()
p.add_run('ก่อน ')
p.add_run(_MARKER + '123' + _MARKER)
p.add_run(' หลัง')
process_runs_to_underline(doc)
for i, r in enumerate(doc.paragraphs[0].runs):
    color = None
    if r.font.color is not None and r.font.color.rgb is not None:
        color = r.font.color.rgb
    print(i, repr(r.text), r.font.underline, color)
